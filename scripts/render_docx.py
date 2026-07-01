#!/usr/bin/env python3
"""
render_docx.py — build a marked-up Word doc from a diff_engine.py pair.

Usage:
    python3 render_docx.py --diffs /tmp/redline-diffs.json --pair v1_v2 --out ./redline-output/

Deletions are shown as red strikethrough, insertions as green underline. This is a
styled-run redline (works in Word, always visible), not native Word Track Changes XML —
that's a stretch goal, not required for v1. See reference/redline-skill-PRD.md.
"""
import argparse
import json
import os
import sys

from docx import Document
from docx.shared import RGBColor


DELETE_COLOR = RGBColor(0xC0, 0x00, 0x00)  # red
INSERT_COLOR = RGBColor(0x00, 0x80, 0x00)  # green


def find_pair(diffs, pair_id):
    for p in diffs["pairs"]:
        if p["pair_id"] == pair_id:
            return p
    return None


def add_ops_to_paragraph(paragraph, ops):
    for op in ops:
        text = op["text"]
        if text == "":
            continue
        run = paragraph.add_run(text)
        if op["type"] == "delete":
            run.font.strike = True
            run.font.color.rgb = DELETE_COLOR
        elif op["type"] == "insert":
            run.font.underline = True
            run.font.color.rgb = INSERT_COLOR
        # 'equal' runs keep default formatting


def render(pair, out_path):
    doc = Document()
    doc.add_heading(f"Redline: {os.path.basename(pair['original'])} \u2192 {os.path.basename(pair['revised'])}", level=1)

    summary = doc.add_paragraph()
    summary.add_run(
        f"{pair['insertions']} insertion(s), {pair['deletions']} deletion(s) across "
        f"{pair['changed_paragraphs']} changed paragraph(s); {pair['unchanged_paragraphs']} unchanged."
    ).italic = True
    doc.add_paragraph()  # spacer

    for para in pair["paragraphs"]:
        p = doc.add_paragraph()
        if para["status"] == "unchanged":
            p.add_run(para["ops"][0]["text"])
        else:
            add_ops_to_paragraph(p, para["ops"])

    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser(description="Render a marked-up docx redline for one version pair.")
    parser.add_argument("--diffs", required=True, help="Path to diff_engine.py's output JSON.")
    parser.add_argument("--pair", required=True, help="pair_id to render (e.g. v1_v2).")
    parser.add_argument("--out", required=True, help="Output directory.")
    args = parser.parse_args()

    with open(args.diffs, "r", encoding="utf-8") as f:
        diffs = json.load(f)

    pair = find_pair(diffs, args.pair)
    if pair is None:
        print(f"Pair '{args.pair}' not found in {args.diffs}.", file=sys.stderr)
        sys.exit(1)
    if pair.get("skipped"):
        print(f"Pair '{args.pair}' was skipped during diffing ({pair.get('reason')}); nothing to render.", file=sys.stderr)
        sys.exit(1)

    os.makedirs(args.out, exist_ok=True)
    out_path = os.path.join(args.out, f"{args.pair}_redline.docx")
    render(pair, out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
